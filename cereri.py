from docx import Document
from tkinter import Checkbutton, IntVar, Toplevel, ttk, messagebox
from tkinter import *
from PIL.Image import Image
from PIL import ImageTk
try:
    from PIL import Image
except ImportError:
    print("no")
from functools import partial

photoimg_ok1=0

tip_doc=[
    "Declarație de primire in spațiu",  #anexa12
    "Declarație de locuire efectivă",   #anexa14
    "Procură specială pentru depunere cerere si ridicarea actului de identitate" #procura
]

title=[
    "DECLARAȚIE",
    "DECLARAȚIE",
    "PROCURĂ SPECIALĂ"
]
    
var=[]
def selecteaza_documente(root):
    global photoimg_ok1,var
    select_frame = Toplevel(root)
    select_frame.title("Verifica informatiile")
    check=[]
    for i in range(len(tip_doc)):
        var.append(IntVar(select_frame,0))
        check.append(Checkbutton(select_frame, text=tip_doc[i], variable=var[i]))
        check[i].grid(row=i)
    img_ok1 = Image.open("icons/id_ok.png")
    img_ok1 = img_ok1.resize((img_ok1.size[0]//2,img_ok1.size[1]//2), Image.ANTIALIAS)
    photoimg_ok1 =  ImageTk.PhotoImage(img_ok1)
    img_nok1 = Image.open("icons/checked.png")
    valid_button=ttk.Button(select_frame,text="", image=photoimg_ok1, compound=LEFT,command=partial(ok_validate,select_frame))
    valid_button.grid(column=0,row=len(tip_doc))

def ok_validate(img_frame):
    ok = messagebox.askokcancel(title="Documente selectate",message="Apasati pe Ok pentru a continua")
    if ok:
        img_frame.destroy()
        img_frame.update()
    else:
        pass

def creaza_cont(informatii):
    continut=[
        """Subsemnatul(a) """+informatii['Nume'].get()+""" """+informatii['Prenume'].get()+""" fiul(fiica)lui şi al ___________________________ \
născut(ă) la data de  """+informatii['Data nastere'].get()+""" în localitatea """+informatii['Loc Nastere'].get()+""" judeţul """+informatii['Seria'].get()+"""\
 posesor al actului de identitate seria """+informatii['Seria'].get()+""" numărul """+informatii['Nr'].get()+"""\
 proprietar al locuinţei din """+informatii['Domiciliu'].get()+""" având actul de spaţiu (denumirea) \
__________________________________nr. _______din _____________emis \
de _______________________________________declar că primesc în spaţiul \
meu de locuit pe _________________________________fiul (fiica) lui \
____________________şi al_____________________________________ \
născut(ă) la data de___________în localitatea_____________________ \
judeţul_________________________________cu ultimul domiciliu în \
localitatea _________________________ str.________________________ \
nr. _____ bl.____ sc._____ et. _____apt._____ sector ________în calitate de \
____________________.
Dau prezenta declaraţie pentru a-i servi numitului(ei) ____________________ \
______________________la schimbarea adresei de domiciliu în spaţiul de \
locuit de la adresa mai sus amintită.
Declar că imobilul a fost/nu a fost notificat în Cartea Funciară ca locuinţă a \
familiei1.""",
        """Subsemnatul(a) """+informatii['Nume'].get()+""" """+informatii['Prenume'].get()+""" fiul(fiica ) \
lui _______________________________ şi al _______________________________ \
născut(ă) la data de """+informatii['Data nastere'].get()+""" în localitatea """+informatii['Loc Nastere'].get()+""" posesor \
al actului de identitate seria """+informatii['Seria'].get()+""" nr. """+informatii['Nr'].get()+""" declar că locuiesc efectiv în \
 """+informatii['Domiciliu'].get()+""" într-un imobil cu destinaţie de locuinţă.
Cele declarate mai sus pot fi confirmate de dl(d-na) _____________________ \
_____________, care locuieşte în_________________________________________ \
_______________________________________________________________________________.
De asemenea, declar că mi-au fost aduse la cunoştinţă prevederile, potrivit cărora, \
falsul în declaraţii constituie infracţiune şi se pedepseşte conform dispoziţiilor \
Codului penal.""",
    """Subsemnatul(a) """+informatii['Nume'].get()+""" """+informatii['Prenume'].get()+""", \
cetăţean (ă) român (ă), născut (ă) la data de """+informatii['Data nastere'].get()+""", \
în localitatea """+informatii['Loc Nastere'].get()+""" fiul(fiica)lui \
_________________________ şi al(a)__________________________________, \
cu domiciliul în """+informatii['Domiciliu'].get()+""" aflat temporar în ________________________________________________, \
identificat(ă) cu paşaport nr. _________________________ în care este înscris \
C.N.P. """+informatii['CNP'].get()+""" cu înălţimea de ________________, ochi \
__________________, semne particulare ____________________________, \
împuternicesc pe ____________________________________________________ \
născut(ă) la data de_______________________________în localitatea \
___________________, identificat(ă) cu paşaport/C.I. seria/nr.___________ CNP \
_______________________________________________ cu domiciliul \
în__________________________ _____________________________________, \
să mă reprezinte în faţa autorităţilor române competente în legătură cu depunerea \
cererii pentru eliberarea actului de identitate şi/sau ridicarea actului de identitate ca \
urmare a1 : _________________________________________________
Pentru aducerea la îndeplinire a prezentului mandat, mandatarul meu va \
semna în numele meu şi pentru mine oriunde va fi necesar, semnătura sa fiindu-mi \
opozabilă.
Redactată şi autentificată astăzi _____________la Ambasada /Consulatul General \
a (al) României la __________________________________"""
    ]
    return continut

def fa_cereri(informatii,var):
    continut=creaza_cont(informatii)
    for i,sel in enumerate(var):
        if sel.get():
            document = Document()
            p_title=document.add_paragraph(title[i])
            p_title.alignment=1
            p=document.add_paragraph('\t'+continut[i])
            document.save(tip_doc[i]+" - completat.docx")
